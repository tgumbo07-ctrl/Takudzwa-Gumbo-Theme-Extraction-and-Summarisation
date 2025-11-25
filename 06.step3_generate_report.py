#!/usr/bin/env python3
"""
STEP 3/3: Loads analysis results and generates the final structured DOCX report.
"""
import pickle
import datetime
from IPython.display import display, HTML
from google.colab.files import download as colab_download

# Try to import necessary libraries, installing them if missing
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    !pip install python-docx
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Input/Output Files ---
INPUT_DATA_FILE = 'analysis_data_step2.pkl'
REPORT_FILENAME = 'Deep_Learning_Analysis_Report.docx'

def create_word_report(data):
    """
    Creates a detailed Word document (.docx) based on the thematic analysis results.
    """
    print(f"\n[Report Generator] Attempting to generate '{REPORT_FILENAME}'...")

    # Unpack data
    executive_summary = data['executive_summary']
    topic_info = data['topic_info']
    total_pages = data['total_pages']
    word_count = data['word_count']
    sentences_analyzed = data['sentences_analyzed']
    generated_timestamp = data['timestamp']

    doc = Document()

    # Set default font size and style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # --- TITLE PAGE ---
    title = doc.add_heading('COMPREHENSIVE DEEP LEARNING PROJECT ANALYSIS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('AI-Generated Thematic and Extractive Summary')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    doc.add_paragraph('\n')  # Add spacing

    # --- METADATA SECTION ---
    doc.add_heading('Analysis Details', level=2)

    p = doc.add_paragraph()
    p.add_run('Source Document: ').bold = True; p.add_run('merged_document_for_analysis.pdf\n')
    p.add_run('Total Pages in Source: ').bold = True; p.add_run(f'{total_pages:,}\n')
    p.add_run('Total Estimated Words: ').bold = True; p.add_run(f'{word_count:,}\n')
    p.add_run('Sentences Analyzed (BERTopic): ').bold = True; p.add_run(f'{sentences_analyzed:,}\n')
    p.add_run('Generated: ').bold = True; p.add_run(generated_timestamp + '\n')
    p.add_run('Method: ').bold = True; p.add_run('BERTopic Neural Topic Modeling (all-MiniLM-L6-v2 embeddings)')

    doc.add_page_break()

    # --- SECTION 1: EXECUTIVE SUMMARY ---
    doc.add_heading('1. Executive Summary', level=1)

    doc.add_paragraph(
        "This report presents an automated, semantic analysis of the merged Deep Learning project material. "
        "The analysis utilized BERTopic to extract themes and representative text, providing a non-biased, "
        "data-driven overview of the document's core content."
    )

    # Add the extractive summary in a blockquote style
    p_summary = doc.add_paragraph()
    p_summary.add_run("Representative Content (Extractive Summary): ").bold = True
    p_summary.add_run(executive_summary)


    # --- SECTION 2: QUANTITATIVE AND THEMATIC EVIDENCE ---
    doc.add_heading('2. Quantitative and Thematic Evidence', level=1)

    # Subtract 1 because BERTopic topic -1 is 'Outlier'
    distinct_themes = len(topic_info[topic_info['Topic'] != -1])
    doc.add_paragraph(f"The BERTopic analysis identified **{distinct_themes}** distinct, non-outlier themes across the document corpus.")

    doc.add_heading('Top 5 Major Themes', level=2)

    # Generate the table for the top 5 themes
    table_data = topic_info[topic_info['Topic'] != -1].head(5)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Theme ID'
    hdr_cells[1].text = 'Count'
    hdr_cells[2].text = 'Keywords & Focus'

    for index, row in table_data.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['Topic'])
        cells[1].text = str(row['Count'])

        # Format theme name and keywords
        theme_content = doc.add_paragraph()
        theme_content.add_run(f"{row['Name'].title()}").bold = True
        theme_content.add_run(f" (Top 3 Keywords: {', '.join(row['Representation'][:3])})")

        # Transfer content to cell[2]
        for paragraph in theme_content.runs:
            cell_run = cells[2].add_paragraph().add_run(paragraph.text)
            cell_run.bold = paragraph.bold

        cells[2].paragraphs[0].clear() # Clear the initial empty paragraph in the cell


    doc.add_paragraph('\n')

    # --- SECTION 3: DETAILED THEMATIC FINDINGS ---
    doc.add_heading('3. Detailed Thematic Breakdown', level=1)
    doc.add_paragraph("The following sections provide a more detailed narrative for the top three themes extracted, derived from the core keywords and representative documents.")

    top_3_themes = topic_info[topic_info['Topic'] != -1].head(3)

    for rank, (index, row) in enumerate(top_3_themes.iterrows()):
        topic_id = row['Topic']
        topic_name = row['Name']
        keywords = ", ".join(row['Representation'])

        doc.add_heading(f"A.{rank+1} Theme {topic_id}: {topic_name.title()}", level=2)

        # Placeholder text derived from keywords
        placeholder_text = (
            f"This theme is highly focused on **{topic_name}**. The key terms, including "
            f"'{keywords}', suggest this section of the document primarily addresses "
            f"the [Describe the core focus - e.g., implementation challenges, evaluation metrics, or ethical considerations] "
            f"related to this area. The representative documents indicate a critical discussion of [Specific finding/challenge]. "
            "**Action Required:** Please review the representative sentences in the analysis log to provide a detailed narrative here."
        )
        doc.add_paragraph(placeholder_text)

    # --- SECTION 4: METHODOLOGY NOTES ---
    doc.add_heading('4. Methodology Notes', level=1)
    doc.add_paragraph(
        "The thematic analysis utilized the **BERTopic** framework. This technique leverages transformer-based embeddings (specifically **all-MiniLM-L6-v2**) to map document segments (sentences) into a semantic space. "
        "It then uses UMAP for dimensionality reduction and HDBSCAN for clustering to identify dense clusters, which represent the themes. "
        "The model employs a c-TF-IDF score to determine the most representative words for each cluster (theme). This analysis specifically avoided common stop words to ensure the identified themes focus on the core research substance."
    )

    # Save the file
    doc.save(REPORT_FILENAME)
    print(f"\n[Report Generator] ✅ Success! Created Word document: {REPORT_FILENAME}")

    # --- Display Download Instructions ---
    download_command = f"colab_download('{REPORT_FILENAME}')"

    html_output = f"""
    <div style="padding: 20px; border: 2px solid #3b82f6; border-radius: 8px; background-color: #e0f2fe; margin-top: 20px;">
        <h3 style="color: #1d4ed8; margin-top: 0;">⬇️ Final Report Generated!</h3>
        <p>Run the code below to download the structured Word document, populated with the BERTopic analysis results.</p>
        <pre style="background-color: #dbeafe; padding: 10px; border-radius: 4px; overflow-x: auto;">{download_command}</pre>
    </div>
    """
    display(HTML(html_output))


# --- Main Report Execution ---
if __name__ == "__main__":
    try:
        # Load data from step 2
        with open(INPUT_DATA_FILE, 'rb') as f:
            report_data = pickle.load(f)

        # Proceed to report generation
        create_word_report(report_data)

    except FileNotFoundError:
        print(f"❌ Input file '{INPUT_DATA_FILE}' not found. Please ensure you have run '01_extraction.py' and '02_analysis.py' successfully.")
    except Exception as e:
        print(f"❌ An error occurred during report generation: {e}")
